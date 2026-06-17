"""
Document Generator — Creates court-ready Word (.docx) and PDF files.
Uses ReportLab for PDF (full Unicode, professional layout) and
docxtpl for Word (Jinja2 template-based with attorney letterhead).
"""

import io
import os
import re
import logging
import json
from datetime import datetime
from typing import Optional, Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.colors import HexColor, black, grey
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, KeepTogether
)
from reportlab.lib import colors


# ═══════════════════════════════════════════════════════════════════════
#  ATTORNEY PROFILE
# ═══════════════════════════════════════════════════════════════════════

class AttorneyProfile:
    def __init__(self, firm_name="", attorney_name="", address="",
                 phone="", email="", bar_number="", state="",
                 law_school="", graduation_year="", years_practice="",
                 specialization="", federal_admission=False, federal_court="",
                 notary_number="", professional_certifications="", 
                 service_address="", fax="", website=""):
        self.firm_name = firm_name
        self.attorney_name = attorney_name
        self.address = address
        self.phone = phone
        self.email = email
        self.bar_number = bar_number
        self.state = state
        self.law_school = law_school
        self.graduation_year = graduation_year
        self.years_practice = years_practice
        self.specialization = specialization
        self.federal_admission = federal_admission
        self.federal_court = federal_court
        self.notary_number = notary_number
        self.professional_certifications = professional_certifications
        self.service_address = service_address
        self.fax = fax
        self.website = website

    @property
    def is_empty(self):
        return not any([self.firm_name, self.attorney_name])

    def to_dict(self):
        return {k: v for k, v in self.__dict__.items()}

    @classmethod
    def from_dict(cls, d):
        return cls(**{k: d.get(k, "") for k in cls.__dict__["__init__"].__code__.co_varnames[1:]})

    def signature_block(self):
        if self.is_empty:
            return ""
        lines = []
        if self.attorney_name:
            lines.append(self.attorney_name)
        if self.firm_name:
            lines.append(self.firm_name)
        if self.bar_number:
            lines.append(f"Bar No. {self.bar_number}")
        if self.state:
            lines.append(f"Admitted to practice in {self.state}")
        if self.federal_admission and self.federal_court:
            lines.append(f"Federal Court: {self.federal_court}")
        return "\n".join(lines)

    def letterhead_lines(self):
        lines = []
        if self.firm_name:
            lines.append(self.firm_name)
        if self.attorney_name:
            lines.append(self.attorney_name)
        if self.address:
            lines.append(self.address)
        if self.service_address and self.service_address != self.address:
            lines.append(f"Service Address: {self.service_address}")
        if self.phone or self.email:
            parts = [p for p in [self.phone, self.email, self.fax] if p]
            lines.append(" | ".join(parts))
        if self.bar_number and self.state:
            lines.append(f"Bar No. {self.bar_number}, {self.state}")
        if self.law_school and self.graduation_year:
            lines.append(f"{self.law_school}, {self.graduation_year}")
        if self.years_practice:
            lines.append(f"{self.years_practice} years of practice")
        if self.specialization:
            lines.append(f"Specialization: {self.specialization}")
        if self.notary_number:
            lines.append(f"Notary Public No. {self.notary_number}")
        if self.professional_certifications:
            lines.append(f"Certifications: {self.professional_certifications}")
        if self.website:
            lines.append(self.website)
        return "\n".join(lines) if lines else ""

    def certification_block(self):
        if self.is_empty:
            return ""
        lines = []
        lines.append("I HEREBY CERTIFY THAT:")
        lines.append("")
        lines.append("1. I am admitted to practice law in the State of " + (self.state or "__________") + ".")
        lines.append("2. I have reviewed the facts and legal authorities cited herein.")
        lines.append("3. The foregoing statements are true and accurate to the best of my knowledge.")
        lines.append("4. This document is filed in accordance with the applicable rules of civil procedure.")
        lines.append("")
        lines.append(self.attorney_name or "")
        lines.append(self.firm_name or "")
        if self.bar_number:
            lines.append(f"Bar No. {self.bar_number}")
        if self.state:
            lines.append(f"State of Admission: {self.state}")
        if self.federal_admission and self.federal_court:
            lines.append(f"Federal Court Admission: {self.federal_court}")
        lines.append(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        return "\n".join(lines)

    def verification_block(self):
        if self.is_empty:
            return ""
        lines = []
        lines.append("VERIFICATION")
        lines.append("=" * 50)
        lines.append("")
        lines.append("I declare under penalty of perjury that:")
        lines.append("")
        lines.append("1. The matters stated in this document are within my personal knowledge;")
        lines.append("2. They are true and accurate;")
        lines.append("3. I have the authority to make this filing;")
        lines.append("4. This document is filed in compliance with all applicable statutes and rules.")
        lines.append("")
        lines.append("Subscribed and sworn to before me this ____ day of _________________, 20____.")
        lines.append("")
        lines.append("Notary Public")
        lines.append(f"Notary Commission No.: {self.notary_number or '__________'}")
        lines.append(f"Commission Expires: _______________, 20______")
        lines.append("")
        lines.append("_________________________________")
        lines.append("Signature of Attorney")
        return "\n".join(lines)

    def get_attorney_credentials(self) -> dict:
        """Return complete attorney credentials for document generation."""
        return {
            "full_name": self.attorney_name or "",
            "firm_name": self.firm_name or "",
            "bar_number": self.bar_number or "",
            "state": self.state or "",
            "law_school": self.law_school or "",
            "graduation_year": self.graduation_year or "",
            "years_practice": self.years_practice or "",
            "specialization": self.specialization or "",
            "federal_admission": self.federal_admission,
            "federal_court": self.federal_court or "",
            "notary_number": self.notary_number or "",
            "professional_certifications": self.professional_certifications or "",
            "address": self.address or "",
            "phone": self.phone or "",
            "email": self.email or "",
            "fax": self.fax or "",
            "website": self.website or "",
            "service_address": self.service_address or "",
        }


# ═══════════════════════════════════════════════════════════════════════
#  COURT METADATA AND VALIDATION
# ═══════════════════════════════════════════════════════════════════════

class CourtMetadata:
    def __init__(self, court_name="", case_number="", filing_date="",
                 judge_name="", court_address="", jurisdiction="",
                 filing_fee="", service_date="", case_type=""):
        self.court_name = court_name
        self.case_number = case_number
        self.filing_date = filing_date
        self.judge_name = judge_name
        self.court_address = court_address
        self.jurisdiction = jurisdiction
        self.filing_fee = filing_fee
        self.service_date = service_date
        self.case_type = case_type

    def to_dict(self):
        return {k: v for k, v in self.__dict__.items()}

    def is_valid(self) -> bool:
        return bool(self.court_name and self.case_number)

    def get_header_text(self) -> str:
        if self.court_name and self.case_number:
            return f"{self.court_name}\nCase No.: {self.case_number}"
        elif self.court_name:
            return self.court_name
        return ""

    def get_footer_text(self) -> str:
        parts = []
        if self.judge_name:
            parts.append(f"Hon. {self.judge_name}")
        if self.filing_date:
            parts.append(f"Filed: {self.filing_date}")
        if parts:
            return " | ".join(parts)
        return ""


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('legal_document_generation.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════
#  ERROR HANDLING AND LOGGING
# ═══════════════════════════════════════════════════════════════════════

class DocumentGenerationError(Exception):
    """Base exception for document generation errors."""
    pass

class ValidationError(DocumentGenerationError):
    """Exception raised for validation errors."""
    pass

class TemplateError(DocumentGenerationError):
    """Exception raised for template-related errors."""
    pass

class IntegrityError(DocumentGenerationError):
    """Exception raised for document integrity errors."""
    pass


def _validate_and_log(operation: str, validation_errors: List[str], 
                     profile: AttorneyProfile = None, metadata: dict = None):
    """Validate inputs and log the operation."""
    if validation_errors:
        error_msg = f"{operation} validation failed: {'; '.join(validation_errors)}"
        logger.error(error_msg)
        raise ValidationError(error_msg)
    
    logger.info(f"{operation} completed successfully")
    if profile and not profile.is_empty:
        logger.info(f"Document generated for {profile.attorney_name} at {profile.firm_name}")
    if metadata and metadata.get("court_name"):
        logger.info(f"Court: {metadata['court_name']}, Case: {metadata.get('case_number', 'N/A')}")


def _safe_operation(operation_name: str, operation_func, *args, **kwargs):
    """Execute an operation safely with error handling and logging."""
    try:
        logger.info(f"Starting {operation_name}")
        result = operation_func(*args, **kwargs)
        logger.info(f"{operation_name} completed successfully")
        return result
    except Exception as e:
        logger.error(f"{operation_name} failed: {str(e)}")
        raise


# ═══════════════════════════════════════════════════════════════════════
#  REPORTLAB STYLES
# ═══════════════════════════════════════════════════════════════════════

LETTER_W, LETTER_H = letter  # 612 x 792 points
MARGIN = inch  # 72 points

def _build_styles():
    ss = getSampleStyleSheet()
    styles = {
        "body": ParagraphStyle(
            "Body",
            parent=ss["Normal"],
            fontName="Times-Roman",
            fontSize=12,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        ),
        "body_indent": ParagraphStyle(
            "BodyIndent",
            parent=ss["Normal"],
            fontName="Times-Roman",
            fontSize=12,
            leading=16,
            alignment=TA_JUSTIFY,
            leftIndent=36,
            spaceAfter=6,
        ),
        "bold": ParagraphStyle(
            "Bold",
            parent=ss["Normal"],
            fontName="Times-Bold",
            fontSize=12,
            leading=16,
            spaceAfter=4,
        ),
        "bold_center": ParagraphStyle(
            "BoldCenter",
            parent=ss["Normal"],
            fontName="Times-Bold",
            fontSize=13,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=8,
            spaceBefore=6,
        ),
        "center": ParagraphStyle(
            "Center",
            parent=ss["Normal"],
            fontName="Times-Roman",
            fontSize=12,
            leading=16,
            alignment=TA_CENTER,
            spaceAfter=4,
        ),
        "right": ParagraphStyle(
            "Right",
            parent=ss["Normal"],
            fontName="Times-Roman",
            fontSize=12,
            alignment=TA_RIGHT,
            spaceAfter=4,
        ),
        "caption": ParagraphStyle(
            "Caption",
            parent=ss["Normal"],
            fontName="Times-Roman",
            fontSize=12,
            leading=16,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "caption_bold": ParagraphStyle(
            "CaptionBold",
            parent=ss["Normal"],
            fontName="Times-Bold",
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=4,
        ),
        "header": ParagraphStyle(
            "Header",
            fontName="Times-Bold",
            fontSize=10,
            alignment=TA_CENTER,
            textColor=grey,
            spaceAfter=2,
        ),
        "footer": ParagraphStyle(
            "Footer",
            fontName="Times-Italic",
            fontSize=9,
            alignment=TA_CENTER,
            textColor=grey,
        ),
        "letterhead": ParagraphStyle(
            "Letterhead",
            fontName="Times-Bold",
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=2,
        ),
        "letterhead_sub": ParagraphStyle(
            "LetterheadSub",
            fontName="Times-Roman",
            fontSize=11,
            alignment=TA_CENTER,
            spaceAfter=1,
        ),
        "letterhead_detail": ParagraphStyle(
            "LetterheadDetail",
            fontName="Times-Roman",
            fontSize=10,
            alignment=TA_CENTER,
            spaceAfter=1,
            textColor=HexColor("#333333"),
        ),
        "signature": ParagraphStyle(
            "Signature",
            fontName="Times-Roman",
            fontSize=12,
            leading=16,
            spaceBefore=16,
            spaceAfter=2,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            fontName="Times-Roman",
            fontSize=12,
            leading=16,
            leftIndent=24,
            spaceAfter=3,
        ),
        "numbered": ParagraphStyle(
            "Numbered",
            fontName="Times-Roman",
            fontSize=12,
            leading=16,
            leftIndent=36,
            spaceAfter=4,
        ),
    }
    return styles

S = _build_styles()


def _validate_document_content(content: str, doc_type: str) -> list:
    """Validate document content for required elements."""
    errors = []
    
    if not content or not content.strip():
        errors.append("Document content is empty")
    
    if doc_type in ("complaint", "motion", "answer"):
        if len(content.strip()) < 100:
            errors.append(f"{doc_type.title()} appears too brief (minimum 100 characters)")
    
    if doc_type == "complaint":
        required_elements = ["plaintiff", "defendant", "cause of action", "relief"]
        for element in required_elements:
            if element.lower() not in content.lower():
                errors.append(f"Complaint missing required element: {element}")
    
    if doc_type == "contract":
        required_elements = ["parties", "consideration", "term", "signature"]
        for element in required_elements:
            if element.lower() not in content.lower():
                errors.append(f"Contract missing required element: {element}")
    
    return errors


def _validate_attorney_profile(profile: AttorneyProfile, doc_type: str = None) -> list:
    """Validate attorney profile for completeness."""
    errors = []
    
    if doc_type in ("complaint", "answer", "motion", "affidavit"):
        if not profile.attorney_name:
            errors.append("Attorney name is required")
        
        if not profile.bar_number:
            errors.append("Bar number is required")
        
        if not profile.state:
            errors.append("State of admission is required")
    
    if profile.law_school and not profile.graduation_year:
        errors.append("Graduation year required when law school is provided")
    
    return errors


def _validate_court_metadata(metadata: CourtMetadata, doc_type: str = None) -> list:
    """Validate court metadata for completeness."""
    errors = []
    
    if doc_type in ("complaint", "answer", "motion", "affidavit"):
        if not metadata.court_name:
            errors.append("Court name is required")
        
        if not metadata.case_number:
            errors.append("Case number is required")
        
        if not metadata.jurisdiction:
            errors.append("Jurisdiction is required")
    
    return errors


def _validate_inputs(content: str, doc_type: str, profile: AttorneyProfile, 
                    metadata: dict) -> List[str]:
    """Validate all inputs for document generation."""
    all_errors = []
    
    # Validate content
    content_errors = _validate_document_content(content, doc_type)
    all_errors.extend(content_errors)
    
    # Validate attorney profile
    profile_errors = _validate_attorney_profile(profile, doc_type)
    all_errors.extend(profile_errors)
    
    # Validate court metadata
    court_metadata = CourtMetadata(
        court_name=metadata.get("court_name", ""),
        case_number=metadata.get("case_number", ""),
        filing_date=metadata.get("filing_date", ""),
        judge_name=metadata.get("judge_name", ""),
        court_address=metadata.get("court_address", ""),
        jurisdiction=metadata.get("jurisdiction", ""),
        filing_fee=metadata.get("filing_fee", ""),
        service_date=metadata.get("service_date", ""),
        case_type=metadata.get("case_type", "")
    )
    court_errors = _validate_court_metadata(court_metadata, doc_type)
    all_errors.extend(court_errors)
    
    return all_errors


# ═══════════════════════════════════════════════════════════════════════
#  UTILITY HELPERS
# ═══════════════════════════════════════════════════════════════════════

def _escape_for_pdf(text: str) -> str:
    """XML-escape text for ReportLab (which uses XML for paragraph content)."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return text


def _cmap(text: str) -> str:
    """Escape text content for PDF flowables."""
    return _escape_for_pdf(text)


def _signature_line(canvas, doc):
    """Draw a signature line on each page (for letter-style docs)."""
    pass


def _add_letterhead(profile: AttorneyProfile, story: list):
    """Add attorney letterhead to the document story."""
    if profile.is_empty:
        return
    lines = []
    if profile.firm_name:
        story.append(Paragraph(_cmap(profile.firm_name), S["letterhead"]))
    if profile.attorney_name:
        story.append(Paragraph(_cmap(profile.attorney_name), S["letterhead_sub"]))
    if profile.address:
        story.append(Paragraph(_cmap(profile.address), S["letterhead_detail"]))
    if profile.phone or profile.email:
        parts = [p for p in [profile.phone, profile.email, profile.fax] if p]
        story.append(Paragraph(_cmap(" | ".join(parts)), S["letterhead_detail"]))
    if profile.bar_number and profile.state:
        story.append(Paragraph(_cmap(f"Bar No. {profile.bar_number}, {profile.state}"), S["letterhead_detail"]))
    if profile.law_school and profile.graduation_year:
        story.append(Paragraph(_cmap(f"{profile.law_school}, {profile.graduation_year}"), S["letterhead_detail"]))
    if profile.years_practice:
        story.append(Paragraph(_cmap(f"{profile.years_practice} years of practice"), S["letterhead_detail"]))
    if profile.specialization:
        story.append(Paragraph(_cmap(f"Specialization: {profile.specialization}"), S["letterhead_detail"]))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#888888")))
    story.append(Spacer(1, 12))


def _add_court_header(metadata: CourtMetadata, story: list):
    """Add court header information to the document story."""
    if not metadata.court_name:
        return
    
    header_text = metadata.get_header_text()
    if header_text:
        story.append(Paragraph(_cmap(header_text), S["header"]))
        story.append(Spacer(1, 4))
    
    if metadata.judge_name:
        story.append(Paragraph(_cmap(f"Hon. {metadata.judge_name}"), S["center"]))
        story.append(Spacer(1, 8))


def _add_court_footer(metadata: CourtMetadata, canvas, doc):
    """Add court footer information to the document."""
    if canvas and doc:
        footer_text = metadata.get_footer_text()
        if footer_text:
            canvas.setFont("Times-Roman", 8)
            canvas.setFillColor(HexColor("#666666"))
            canvas.drawCentredString(w / 2, 0.3 * inch, footer_text)
            canvas.setFillColor(black)


# ═══════════════════════════════════════════════════════════════════════
#  WORD GENERATION (docxtpl)
# ═══════════════════════════════════════════════════════════════════════

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates", "docx")


def _get_template_path(doc_type: str) -> str:
    if doc_type in ("complaint", "motion", "answer", "affidavit"):
        return os.path.join(TEMPLATE_DIR, "court.docx")
    elif doc_type in ("nda", "contract"):
        return os.path.join(TEMPLATE_DIR, "contract.docx")
    return os.path.join(TEMPLATE_DIR, "letter.docx")


def create_word_document(content: str, doc_type: str, metadata: dict,
                         profile: AttorneyProfile = None) -> io.BytesIO:
    """
    Generate a professional Word document using docxtpl with attorney letterhead.
    Enhanced with court metadata, validation, and professional formatting.
    """
    if profile is None:
        profile = AttorneyProfile()

    # Validate all inputs
    validation_errors = _validate_inputs(content, doc_type, profile, metadata)
    _validate_and_log(
        f"Word document generation for {doc_type}",
        validation_errors,
        profile,
        metadata
    )

    template_path = _get_template_path(doc_type)

    if os.path.exists(template_path):
        return _create_with_template(content, doc_type, profile, template_path, metadata)
    return _create_plain_word(content, doc_type, profile, metadata)


def _create_with_template(content: str, doc_type: str,
                          profile: AttorneyProfile, template_path: str,
                          metadata: dict) -> io.BytesIO:
    """Use docxtpl Jinja2 template for Word generation with enhanced court formatting."""
    sig = profile.signature_block()
    sig_text = f"\n{sig}" if sig else ""

    # Prepare enhanced context with court metadata
    context = {
        "firm_name": profile.firm_name or "",
        "attorney_name": profile.attorney_name or "",
        "address": profile.address or "",
        "phone": profile.phone or "",
        "email": profile.email or "",
        "bar_number": profile.bar_number or "",
        "state": profile.state or "",
        "law_school": profile.law_school or "",
        "graduation_year": profile.graduation_year or "",
        "years_practice": profile.years_practice or "",
        "specialization": profile.specialization or "",
        "federal_admission": profile.federal_admission,
        "federal_court": profile.federal_court or "",
        "notary_number": profile.notary_number or "",
        "professional_certifications": profile.professional_certifications or "",
        "service_address": profile.service_address or "",
        "fax": profile.fax or "",
        "website": profile.website or "",
        "body": content,
        "signature_block": sig_text,
        "certification_block": profile.certification_block() or "",
        
        # Court metadata
        "court_name": metadata.get("court_name", ""),
        "case_number": metadata.get("case_number", ""),
        "filing_date": metadata.get("filing_date", ""),
        "judge_name": metadata.get("judge_name", ""),
        "court_address": metadata.get("court_address", ""),
        "jurisdiction": metadata.get("jurisdiction", ""),
        "filing_fee": metadata.get("filing_fee", ""),
        "service_date": metadata.get("service_date", ""),
        "case_type": metadata.get("case_type", ""),
        
        # Document metadata
        "document_date": datetime.now().strftime("%B %d, %Y"),
        "document_type_label": doc_type.replace("_", " ").title(),
    }

    doc = DocxTemplate(template_path)
    doc.render(context)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _create_plain_word(content: str, doc_type: str, 
                       profile: AttorneyProfile = None,
                       metadata: dict = None) -> io.BytesIO:
    """Fallback Word generation without template with enhanced formatting."""
    if profile is None:
        profile = AttorneyProfile()
    if metadata is None:
        metadata = {}
    
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1.5)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    lines = content.split("\n")
    for line in lines:
        line = line.rstrip()
        if not line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            continue

        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        stripped = line.strip()
        if stripped.isupper() and len(stripped) > 3:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif re.match(r"^\s*(v\.|vs\.)\s*$", stripped, re.IGNORECASE):
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif re.match(r"^Case No\.", stripped, re.IGNORECASE):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif re.match(r"^(COUNT|CAUSE OF ACTION|PRAYER|WHEREFORE|JURY DEMAND|ARTICLE|SECTION|RECITALS)", stripped, re.IGNORECASE):
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("____") or line.startswith("___"):
            p.paragraph_format.space_before = Pt(24)

        if doc_type in ("complaint", "motion", "answer", "affidavit"):
            p.paragraph_format.line_spacing = Pt(24)
        else:
            p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.space_after = Pt(2)

    # Add court header
    if metadata.get("court_name") or metadata.get("case_number"):
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if metadata.get("court_name"):
            header_run = header_para.add_run(metadata["court_name"])
            header_run.bold = True
            header_run.font.size = Pt(14)
        if metadata.get("case_number"):
            header_para.add_run(f"\nCase No.: {metadata['case_number']}")
        header_para.paragraph_format.space_after = Pt(12)

    # Add attorney letterhead
    if not profile.is_empty:
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if profile.firm_name:
            header_run = header_para.add_run(profile.firm_name)
            header_run.bold = True
            header_run.font.size = Pt(16)
            header_para.add_run("\n")
        if profile.attorney_name:
            header_para.add_run(profile.attorney_name)
            header_para.add_run("\n")
        if profile.address:
            header_para.add_run(profile.address)
            header_para.add_run("\n")
        if profile.phone or profile.email:
            contact_parts = []
            if profile.phone:
                contact_parts.append(profile.phone)
            if profile.email:
                contact_parts.append(profile.email)
            if profile.fax:
                contact_parts.append(f"Fax: {profile.fax}")
            header_para.add_run(" | ".join(contact_parts))
            header_para.add_run("\n")
        if profile.bar_number and profile.state:
            header_para.add_run(f"Bar No. {profile.bar_number}, {profile.state}")
            header_para.add_run("\n")
        header_para.paragraph_format.space_after = Pt(12)

    # Add document metadata
    metadata_para = doc.add_paragraph()
    metadata_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    metadata_para.style.font.size = Pt(10)
    metadata_para.style.font.italic = True
    
    metadata_lines = []
    if metadata.get("filing_date"):
        metadata_lines.append(f"Filed: {metadata['filing_date']}")
    if metadata.get("judge_name"):
        metadata_lines.append(f"Hon. {metadata['judge_name']}")
    if metadata.get("jurisdiction"):
        metadata_lines.append(f"Jurisdiction: {metadata['jurisdiction']}")
    if metadata_lines:
        metadata_para.add_run(" | ".join(metadata_lines))
    
    metadata_para.paragraph_format.space_after = Pt(24)

    # Add certification block
    if not profile.is_empty:
        cert_para = doc.add_paragraph()
        cert_para.paragraph_format.space_before = Pt(24)
        cert_para.paragraph_format.line_spacing = Pt(16)
        
        cert_text = profile.certification_block()
        if cert_text:
            for cert_line in cert_text.split("\n"):
                cert_para.add_run(cert_line + "\n")
        
        cert_para.paragraph_format.space_after = Pt(24)

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    from docx.oxml.ns import qn
    fc1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fc1)
    r2 = fp.add_run()
    instr = r2._element.makeelement(qn("w:instrText"), {})
    instr.text = " PAGE "
    r2._element.append(instr)
    r3 = fp.add_run()
    fc2 = r3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    r3._element.append(fc2)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  PDF GENERATION (ReportLab)
# ═══════════════════════════════════════════════════════════════════════

def _make_pdf_base(profile: AttorneyProfile = None, metadata: CourtMetadata = None):
    """Create a base SimpleDocTemplate with page callbacks."""
    if profile is None:
        profile = AttorneyProfile()
    buffer = io.BytesIO()

    def on_first_page(canvas, doc):
        canvas.saveState()
        _draw_page_frame(canvas, doc)
        _draw_header_footer(canvas, doc, profile, metadata)
        canvas.restoreState()

    def on_later_pages(canvas, doc):
        canvas.saveState()
        _draw_page_frame(canvas, doc)
        _draw_header_footer(canvas, doc, profile, metadata)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=MARGIN,
        bottomMargin=MARGIN,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
    )
    return buffer, doc, on_first_page, on_later_pages


def _draw_page_frame(canvas, doc):
    """Draw a thin border frame on each page."""
    w, h = letter
    margin = 0.5 * inch
    canvas.setStrokeColor(HexColor("#CCCCCC"))
    canvas.setLineWidth(0.3)
    canvas.rect(margin, margin, w - 2 * margin, h - 2 * margin)


def _draw_header_footer(canvas, doc, profile: AttorneyProfile, metadata: CourtMetadata = None):
    """Draw header line, court information, and page number."""
    w, h = letter
    
    # Header line
    canvas.setStrokeColor(HexColor("#999999"))
    canvas.setLineWidth(0.4)
    y_header = h - 0.6 * inch
    canvas.line(0.75 * inch, y_header, w - 0.75 * inch, y_header)

    # Court information (left side)
    if metadata:
        canvas.setFont("Times-Roman", 9)
        canvas.setFillColor(HexColor("#333333"))
        y_court = y_header - 0.15
        court_lines = []
        if metadata.court_name:
            court_lines.append(metadata.court_name)
        if metadata.judge_name:
            court_lines.append(f"Hon. {metadata.judge_name}")
        if metadata.jurisdiction:
            court_lines.append(metadata.jurisdiction)
        
        for i, line in enumerate(court_lines):
            canvas.drawString(0.75 * inch, y_court - (i * 12), line)
    
    # Footer with page number and case info
    canvas.setFont("Times-Roman", 9)
    canvas.setFillColor(grey)
    
    # Page number (right side)
    canvas.drawRightString(w - 0.75 * inch, 0.45 * inch, f"Page {doc.page}")
    
    # Case number (left side)
    if metadata and metadata.case_number:
        canvas.drawString(0.75 * inch, 0.45 * inch, f"Case No.: {metadata.case_number}")
    
    canvas.setFillColor(black)


def _parse_blocks(content: str):
    """Split content into paragraph blocks for PDF layout."""
    lines = content.split("\n")
    blocks = []
    current = []
    for ln in lines:
        s = ln.rstrip()
        if s == "":
            if current:
                blocks.append("\n".join(current))
                current = []
        else:
            current.append(s)
    if current:
        blocks.append("\n".join(current))
    return blocks


def _is_centered(text: str) -> bool:
    t = text.strip()
    if t.isupper() and len(t) > 3:
        return True
    if re.match(r"^\s*(v\.|vs\.)\s*$", t, re.IGNORECASE):
        return True
    if re.match(r"^Case No\.", t, re.IGNORECASE):
        return True
    if re.match(r"^IN THE ", t, re.IGNORECASE):
        return True
    if re.match(r"^UNITED STATES", t, re.IGNORECASE):
        return True
    return False


def _is_bold_heading(text: str) -> bool:
    t = text.strip()
    if re.match(r"^(COUNT|CAUSE OF ACTION|PRAYER|WHEREFORE|JURY DEMAND|ARTICLE|SECTION|RECITALS)", t, re.IGNORECASE):
        return True
    if re.match(r"^(NOW, THEREFORE|IN WITNESS WHEREOF)", t, re.IGNORECASE):
        return True
    return False


def _is_section_header(text: str) -> bool:
    t = text.strip()
    if t.isupper() and len(t) > 3 and not t.startswith("("):
        return True
    if re.match(r"^[A-Z][A-Z\s&/]+:$", t):
        return True
    return False


def create_pdf_document(content: str, doc_type: str, metadata: dict,
                        profile: AttorneyProfile = None) -> io.BytesIO:
    """
    Generate a professional PDF using ReportLab with full Unicode support.
    Enhanced with court metadata, validation, and professional formatting.
    """
    if profile is None:
        profile = AttorneyProfile()
    
    # Validate all inputs
    validation_errors = _validate_inputs(content, doc_type, profile, metadata)
    _validate_and_log(
        f"PDF generation for {doc_type}",
        validation_errors,
        profile,
        metadata
    )
    
    # Extract court metadata from the metadata dict
    court_metadata = CourtMetadata(
        court_name=metadata.get("court_name", ""),
        case_number=metadata.get("case_number", ""),
        filing_date=metadata.get("filing_date", ""),
        judge_name=metadata.get("judge_name", ""),
        court_address=metadata.get("court_address", ""),
        jurisdiction=metadata.get("jurisdiction", ""),
        filing_fee=metadata.get("filing_fee", ""),
        service_date=metadata.get("service_date", ""),
        case_type=metadata.get("case_type", "")
    )
    
    buffer, doc, on_first_page, on_later_pages = _make_pdf_base(profile, court_metadata)
    story = []
    ew = LETTER_W - 2 * MARGIN  # effective width

    # Attorney letterhead
    _add_letterhead(profile, story)
    
    # Court header
    _add_court_header(court_metadata, story)

    blocks = _parse_blocks(content)
    for block in blocks:
        block_stripped = block.strip()
        if not block_stripped:
            story.append(Spacer(1, 6))
            continue

        block_text = block_stripped.replace("\n", " ")

        # Signature lines
        if block_text.startswith("____") or block_text.startswith("___"):
            story.append(Spacer(1, 16))
            continue

        # Centered caption/court lines
        if _is_centered(block_text):
            story.append(Paragraph(_cmap(block_text), S["caption_bold"] if block_text.isupper() else S["center"]))
            continue

        # Bold heading
        if _is_bold_heading(block_text):
            story.append(Spacer(1, 6))
            story.append(Paragraph(_cmap(block_text), S["bold_center"]))
            story.append(Spacer(1, 4))
            continue

        # Section header (all caps)
        if _is_section_header(block_text):
            story.append(Paragraph(_cmap(block_text), S["bold"]))
            story.append(Spacer(1, 2))
            continue

        # Bullet items
        if re.match(r"^[•\-*]\s", block_text):
            clean = re.sub(r"^[•\-*]\s*", "", block_text)
            story.append(Paragraph(_cmap(f"&bull; {clean}"), S["bullet"]))
            continue

        # Numbered paragraphs
        num_match = re.match(r"^(\d+)\.\s+(.+)", block_text)
        if num_match:
            story.append(Paragraph(_cmap(block_text), S["numbered"]))
            continue

        # Sub-items (a), (i), etc.
        if re.match(r"^\s*[\(\[]?[a-zivx]+[\)\]]\s", block_text):
            story.append(Paragraph(_cmap(block_text), S["body_indent"]))
            continue

        # WHEREAS / THEREFORE — italic or bold
        if block_text.upper().startswith("WHEREAS"):
            story.append(Paragraph(_cmap(block_text), ParagraphStyle(
                "Whereas", fontName="Times-Italic", fontSize=12, leading=16,
                leftIndent=24, spaceAfter=6)))
            continue

        if block_text.upper().startswith("NOW, THEREFORE") or block_text.upper().startswith("WHEREFORE"):
            story.append(Spacer(1, 4))
            story.append(Paragraph(_cmap(block_text), S["bold"]))
            story.append(Spacer(1, 4))
            continue

        # Regular paragraph
        story.append(Paragraph(_cmap(block_text), S["body"]))

    # Professional certification block
    if not profile.is_empty:
        story.append(Spacer(1, 24))
        story.append(HRFlowable(width="60%", thickness=0.7, color=HexColor("#444444")))
        story.append(Spacer(1, 8))
        
        # Certification text
        cert_lines = [
            "I HEREBY CERTIFY THAT:",
            "",
            "1. I am admitted to practice law in the State of " + (profile.state or "__________") + ".",
            "2. I have reviewed the facts and legal authorities cited herein.",
            "3. The foregoing statements are true and accurate to the best of my knowledge.",
            "4. This document is filed in accordance with the applicable rules of civil procedure.",
            "",
            profile.signature_block(),
            "",
            f"Attorney at Law - Admitted {datetime.now().year if profile.years_practice else '__________'},",
            f"Bar No. {profile.bar_number or '__________'}, {profile.state or '__________'}"
        ]
        
        for line in cert_lines:
            if line.startswith("1.") or line.startswith("2.") or line.startswith("3.") or line.startswith("4."):
                story.append(Paragraph(_cmap(line), S["body_indent"]))
            elif line.strip():
                story.append(Paragraph(_cmap(line), S["body"]))
            else:
                story.append(Spacer(1, 4))

    doc.build(story, onFirstPage=on_first_page, onLaterPages=on_later_pages)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════════════════════
#  DOCUMENT INTEGRITY AND SECURITY
# ═══════════════════════════════════════════════════════════════════════

import hashlib
import hmac
import json

class DocumentIntegrity:
    def __init__(self, secret_key: str = None):
        self.secret_key = secret_key or "default_secret_key_change_in_production"
        self.document_signatures = {}

    def generate_document_hash(self, content: str, metadata: dict) -> str:
        """Generate SHA-256 hash of document content for integrity verification."""
        hash_data = {
            "content": content,
            "timestamp": datetime.now().isoformat(),
            "metadata": metadata
        }
        json_str = json.dumps(hash_data, sort_keys=True, separators=(',', ':'))
        return hashlib.sha256(json_str.encode()).hexdigest()

    def verify_document_integrity(self, content: str, expected_hash: str, metadata: dict) -> bool:
        """Verify document integrity using stored hash."""
        current_hash = self.generate_document_hash(content, metadata)
        return hmac.compare_digest(current_hash, expected_hash)

    def sign_document(self, content: str, metadata: dict) -> dict:
        """Create a digital signature for the document."""
        document_hash = self.generate_document_hash(content, metadata)
        signature_data = {
            "document_hash": document_hash,
            "timestamp": datetime.now().isoformat(),
            "metadata": metadata,
            "signature": hmac.new(
                self.secret_key.encode(),
                document_hash.encode(),
                hashlib.sha256
            ).hexdigest()
        }
        return signature_data

    def verify_document_signature(self, signature_data: dict, content: str) -> bool:
        """Verify document signature."""
        if "document_hash" not in signature_data or "signature" not in signature_data:
            return False
        
        expected_hash = signature_data["document_hash"]
        current_hash = self.generate_document_hash(content, signature_data.get("metadata", {}))
        
        if not hmac.compare_digest(expected_hash, current_hash):
            return False
        
        expected_signature = hmac.new(
            self.secret_key.encode(),
            expected_hash.encode(),
            hashlib.sha256
        ).hexdigest()
        
        return hmac.compare_digest(expected_signature, signature_data["signature"])


def _generate_document_id() -> str:
    """Generate a unique document ID."""
    return hashlib.sha256(
        f"{datetime.now().isoformat()}{os.urandom(8).hex()}".encode()
    ).hexdigest()[:16]


def _log_document_generation(doc_id: str, doc_type: str, profile: AttorneyProfile,
                           metadata: dict, success: bool, error: str = None):
    """Log document generation for audit purposes."""
    log_entry = {
        "timestamp": datetime.now().isoformat(),
        "document_id": doc_id,
        "document_type": doc_type,
        "attorney_name": profile.attorney_name,
        "firm_name": profile.firm_name,
        "bar_number": profile.bar_number,
        "court_name": metadata.get("court_name", ""),
        "case_number": metadata.get("case_number", ""),
        "success": success,
        "error": error
    }
    
    # In production, this would be written to a secure log file or database
    # For now, we'll just print it for demonstration
    if success:
        print(f"✓ Document generated successfully: {doc_id} ({doc_type})")
    else:
        print(f"✗ Document generation failed: {doc_id} ({doc_type}) - {error}")


def _validate_document_content_for_tampering(content: str, original_content_hash: str) -> bool:
    """Check if document content has been tampered with."""
    current_hash = hashlib.sha256(content.encode()).hexdigest()
    return hmac.compare_digest(current_hash, original_content_hash)


# ═══════════════════════════════════════════════════════════════════════
#  FILENAME HELPER
# ═══════════════════════════════════════════════════════════════════════

def make_filename(doc_type: str, party_name: str = "") -> str:
    date_str = datetime.now().strftime("%Y%m%d")
    clean_name = re.sub(r"[^a-zA-Z0-9]", "_", party_name)[:30] if party_name else "document"
    type_label = doc_type.replace("_", "-")
    return f"{type_label}_{clean_name}_{date_str}"


def create_document_with_integrity(content: str, doc_type: str, metadata: dict,
                                   profile: AttorneyProfile = None,
                                   secret_key: str = None) -> tuple:
    """
    Create a document with integrity and security features.
    Returns a tuple of (document_bytes, document_id, integrity_info).
    """
    if profile is None:
        profile = AttorneyProfile()
    
    # Validate all inputs
    validation_errors = _validate_inputs(content, doc_type, profile, metadata)
    _validate_and_log(
        f"Document generation with integrity for {doc_type}",
        validation_errors,
        profile,
        metadata
    )
    
    # Generate document ID
    doc_id = _generate_document_id()
    
    # Create integrity checker
    integrity_checker = DocumentIntegrity(secret_key)
    
    # Generate document hash
    document_hash = integrity_checker.generate_document_hash(content, metadata)
    
    # Create digital signature
    signature_data = integrity_checker.sign_document(content, metadata)
    
    # Generate document based on type
    if doc_type in ("complaint", "motion", "answer", "affidavit"):
        buffer = create_pdf_document(content, doc_type, metadata, profile)
    else:
        buffer = create_word_document(content, doc_type, metadata, profile)
    
    # Log the generation
    _log_document_generation(doc_id, doc_type, profile, metadata, True)
    
    # Return document with integrity information
    integrity_info = {
        "document_id": doc_id,
        "document_hash": document_hash,
        "signature_data": signature_data,
        "generated_at": datetime.now().isoformat(),
        "profile_hash": hashlib.sha256(json.dumps(profile.to_dict(), sort_keys=True).encode()).hexdigest(),
    }
    
    return buffer, doc_id, integrity_info
